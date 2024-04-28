''' word from python '''
import asyncio
import time
import os

from docx import Document
from docx.shared import Cm, Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

from func.style import setStyle

async def writingDocx(data):
    document = Document()
    if (os.path.exists('demo.docx')): os.remove('demo.docx')
    if (os.path.exists('demo.pdf')): os.remove('demo.pdf')

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(0.5)
        section.right_margin = Cm(0.5)
    
    setStyle(document=document)

    table_comp_name = document.add_table(rows=1, cols=1)
    table_comp_name.style = 'Table Grid'
    table_comp_name.autofit = False
    table_comp_name.allow_autofit = False
    for row in table_comp_name.rows:
        row.height = Cm(0.7)
        for id, cell in enumerate(row.cells):
            cell.width = Cm(14)
    comp_name_row = table_comp_name.rows[0].cells
    comp_name_row[0].text = data['company']['name']
    comp_name_row[0].paragraphs[0].runs[0].font.bold = True
    comp_name_row[0].paragraphs[0].runs[0].font.size = Pt(16)
    comp_name_row[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    table_section2 = document.add_table(rows=1, cols=2)
    table_section2.style = 'Table Grid'
    table_section2.autofit = False
    table_section2.allow_autofit = False
    for row in table_section2.rows:
        row.height = Cm(0.7)
        for id, cell in enumerate(row.cells):
            if id == 1:
                cell.width = Cm(7)
            else:
                cell.width = Cm(14)
    for i, data in enumerate(data['company']['section02'], start=0):
        table_section2.cell(i, 0).text = data[0]
        table_section2.cell(i, 1).text = data[1]
        table_section2.cell(i, 1).paragraphs[0].runs[0].font.bold = True
        table_section2.cell(i, 1).paragraphs[0].runs[0].font.size = Pt(16)
        table_section2.cell(i, 1).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table_section2.cell(i, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    document.save('demo.docx')

    from docx2pdf import convert
    docx_file = 'demo.docx'
    pdf_file = 'demo.pdf'
    convert(docx_file, pdf_file)

def main():
    t1 = time.time()
    data = {
        'company': {
            'name': 'Example company name',
            'section02': [
                ('123/123\tMoo 0\tSubdistrict\tDistrict\tProvince\t12345\n0123-456-7890, 001-002-0023\nTax identification number\t01234567890123\tHead office', 'Quotation'),
            ]
        },
        'customer': {
            
        }
        
    }
    asyncio.run(writingDocx(data))
    t2 = time.time() - t1
    print(f'Executed in {t2:0.2f} seconds.')

if __name__ == '__main__':
    main()